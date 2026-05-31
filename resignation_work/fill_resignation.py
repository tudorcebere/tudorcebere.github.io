from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


SRC = Path("/Users/tudorcebere/Downloads/Modèle démission contrat.docx")
OUT = Path("/Users/tudorcebere/Work/website/tudorcebere.github.io/resignation_work/demission-cebere-ioan-tudor-inria.docx")


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            paragraph._p.remove(child)


def add(paragraph, text="", bold=False, italic=False, line_break=False):
    run = paragraph.add_run()
    if line_break:
        run.add_break()
    run.text = text
    run.bold = bold
    run.italic = italic
    return run


doc = Document(SRC)
p = doc.paragraphs

for paragraph in p:
    clear_paragraph(paragraph)

p[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
add(p[0], "Ioan-Tudor CEBERE", bold=True)
add(p[0], "140, rue Toiras 14b", line_break=True)
add(p[0], "34000 Montpellier", line_break=True)

p[2].alignment = WD_ALIGN_PARAGRAPH.RIGHT
add(p[2], "INRIA", bold=True)
add(p[2], "À l’attention de Monsieur le responsable de l’équipe PREMEDICAL", line_break=True)
add(p[2], "Centre Inria d’Université Côte d’Azur - Antenne de Montpellier", line_break=True)

add(p[3], "Objet : Démission du contrat 82889", bold=True)

add(p[4], "Monsieur,")

add(p[5], "Par la présente, je vous informe de ma décision de démissionner de ")
add(p[5], "mes fonctions de doctorant contractuel", bold=True)
add(p[5], ", exercées depuis le ")
add(p[5], "24 octobre 2022", bold=True)
add(p[5], " au sein de l’INRIA.")

add(p[6], "Conformément aux dispositions de mon contrat de travail, je vous informe que mon ")
add(p[6], "dernier jour de travail sera fixé au 31 août 2026, au soir", bold=True)
add(p[6], ", mettant ainsi fin à mon contrat de travail à cette date.")

add(p[7], "Je vous remercie de bien vouloir prendre acte de cette démission et de me confirmer la date effective de fin de contrat.")

add(p[8], "Lors de mon dernier jour de travail, je vous remercie par ailleurs de bien vouloir me transmettre les documents de fin de contrat, à savoir :")

add(p[9], "le certificat de travail,")
add(p[10], "l’attestation France Travail.")

add(p[11], "Je vous prie d’agréer, Monsieur, l’expression de mes salutations distinguées.")

add(p[12], "Fait à Montpellier, le 31 mai 2026")

add(p[13], "Signature")
add(p[13], "Ioan-Tudor CEBERE", line_break=True)

doc.save(OUT)
print(OUT)
